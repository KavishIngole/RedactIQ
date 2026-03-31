"""Write redacted content back to the original file format.

Supports: DOCX, XLSX, PDF, CSV, and plain text.
Each writer preserves the original structure as much as possible.
"""

from __future__ import annotations

import csv
import io
import copy
from pathlib import Path
from typing import Any

from loguru import logger


def write_redacted_docx(
    original_bytes: bytes,
    redacted_pages: list[dict[str, str]],
    output_path: str,
) -> None:
    """Write redacted content back to a DOCX file, preserving formatting.

    Strategy: walk through each paragraph and table cell in the original doc,
    match to the corresponding redacted text, and replace the runs' text
    while keeping font, size, bold/italic styling intact.
    """
    from docx import Document

    doc = Document(io.BytesIO(original_bytes))

    # Build a mapping of original text -> redacted text from pages
    redact_map: dict[str, str] = {}
    for page in redacted_pages:
        original = page.get("original_text", "")
        redacted = page.get("redacted_text", "")
        if original and redacted and original != redacted:
            redact_map[original.strip()] = redacted.strip()

    # Redact paragraphs
    for para in doc.paragraphs:
        original_text = para.text.strip()
        if original_text in redact_map:
            _replace_paragraph_text(para, redact_map[original_text])

    # Redact table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original_text = para.text.strip()
                    if original_text in redact_map:
                        _replace_paragraph_text(para, redact_map[original_text])

    doc.save(output_path)


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace paragraph text while trying to keep the first run's formatting."""
    if not paragraph.runs:
        paragraph.text = new_text
        return

    # Preserve the first run's formatting, clear the rest
    first_run = paragraph.runs[0]
    for run in paragraph.runs[1:]:
        run.text = ""
    first_run.text = new_text


def write_redacted_xlsx(
    original_bytes: bytes,
    redacted_pages: list[dict[str, str]],
    output_path: str,
) -> None:
    """Write redacted content back to an XLSX file, preserving formatting.

    Strategy: walk through each cell, find its text in the redaction map,
    and replace the value while keeping formatting intact.
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(original_bytes))

    # Build redact map
    redact_map: dict[str, str] = {}
    for page in redacted_pages:
        original = page.get("original_text", "")
        redacted = page.get("redacted_text", "")
        if original and redacted and original != redacted:
            redact_map[original.strip()] = redacted.strip()

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    cell_text = str(cell.value).strip()
                    if cell_text in redact_map:
                        cell.value = redact_map[cell_text]

    wb.save(output_path)


def write_redacted_pdf(
    redacted_pages: list[dict[str, str]],
    output_path: str,
) -> None:
    """Write redacted content to a new PDF file.

    Creates a clean PDF with one page per original page/section.
    Original PDF layout cannot be perfectly preserved, but text content
    and page structure are maintained.
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    for page_info in redacted_pages:
        pdf.add_page()
        pdf.set_font("Helvetica", size=10)

        label = page_info.get("label", "")
        text = page_info.get("redacted_text", "")

        if label:
            pdf.set_font("Helvetica", "B", size=12)
            pdf.cell(0, 10, label, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=10)
            pdf.ln(2)

        # Write text line by line to handle long content
        for line in text.split("\n"):
            # Replace any non-latin1 chars for fpdf compatibility
            safe_line = line.encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 5, safe_line)

    pdf.output(output_path)


def write_redacted_csv(
    redacted_rows: list[list[str]],
    output_path: str,
) -> None:
    """Write redacted CSV rows to a file."""
    with open(output_path, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(redacted_rows)


def write_redacted_text(
    content: str,
    output_path: str,
) -> None:
    """Write redacted text content to a file."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
