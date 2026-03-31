"""File parsing utilities for extracting text from various document formats.

Supports: PDF, DOCX, XLSX, CSV, and plain text files.
"""

from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import Any


def extract_text(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Extract text content from a file based on its extension.

    Args:
        file_bytes: Raw file bytes.
        filename: Original filename (used to detect format).

    Returns:
        Dict with keys:
        - "format": detected format (pdf, docx, xlsx, csv, text)
        - "pages": list of dicts, each with "page" (label) and "text"
          For CSV/XLSX: each page = sheet/table.
          For PDF: each page = physical page.
          For DOCX/text: single page with all content.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    elif ext in (".xlsx", ".xls"):
        return _parse_xlsx(file_bytes)
    elif ext == ".csv":
        return _parse_csv(file_bytes)
    else:
        return _parse_text(file_bytes, filename)


def _parse_pdf(file_bytes: bytes) -> dict[str, Any]:
    """Extract text from PDF, page by page."""
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append({"page": f"Page {i + 1}", "text": text})

    return {"format": "pdf", "pages": pages}


def _parse_docx(file_bytes: bytes) -> dict[str, Any]:
    """Extract text from DOCX (paragraphs + tables)."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))

    # Extract paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    body_text = "\n".join(paragraphs)

    pages = []
    if body_text.strip():
        pages.append({"page": "Body", "text": body_text})

    # Extract tables
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            pages.append({"page": f"Table {i + 1}", "text": table_text})

    return {"format": "docx", "pages": pages}


def _parse_xlsx(file_bytes: bytes) -> dict[str, Any]:
    """Extract text from Excel, sheet by sheet, cell by cell."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    pages = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cell_texts = [str(c) if c is not None else "" for c in row]
            if any(t.strip() for t in cell_texts):
                rows.append(" | ".join(cell_texts))
        if rows:
            pages.append({"page": f"Sheet: {sheet_name}", "text": "\n".join(rows)})

    wb.close()
    return {"format": "xlsx", "pages": pages}


def _parse_csv(file_bytes: bytes) -> dict[str, Any]:
    """Parse CSV into structured rows. Returns raw CSV text as a single page."""
    try:
        content = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        content = file_bytes.decode("latin-1")

    return {"format": "csv", "pages": [{"page": "CSV", "text": content}]}


def _parse_text(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Parse plain text files (TXT, LOG, JSON, MD, etc.)."""
    try:
        content = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        content = file_bytes.decode("latin-1")

    ext = Path(filename).suffix.lower().lstrip(".") or "text"
    return {"format": ext, "pages": [{"page": "Content", "text": content}]}
